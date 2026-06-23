"""
title: Word Document Creator
author: open-webui-aws-fargate
version: 2.0.0
description: Generate a Microsoft Word (.docx) document from a Markdown string. Headings, paragraphs, bullet and numbered lists, and bold/italic inline formatting are supported. The file is delivered as a chat attachment.
required_open_webui_version: 0.9.0
"""

import io
import re
import uuid
from typing import Optional

from pydantic import BaseModel, Field

from docx import Document
from docx.shared import Pt


class Tools:
    class Valves(BaseModel):
        max_size_mb: int = Field(default=10, description="Refuse to attach files larger than this many megabytes.")

    def __init__(self):
        self.valves = self.Valves()
        self.citation = False

    async def create_word_document(
        self,
        title: str,
        markdown: str,
        __event_emitter__=None,
        __user__: Optional[dict] = None,
    ) -> str:
        """
        Create a Microsoft Word (.docx) document and attach it to the chat.

        Args:
            title: Document title. Rendered as the top-level heading and used as the filename.
            markdown: Document body as Markdown text. Supported syntax:
                - Headings: # H1  ## H2  ### H3
                - Paragraphs: plain text separated by blank lines
                - Bullet lists: lines starting with - or *
                - Numbered lists: lines starting with 1. or 1)
                - Bold: **text**   Italic: *text*   Inline code: `text`
        """
        doc = Document()
        doc.add_heading(title, level=0)

        for block in _parse_blocks(markdown):
            kind = block["kind"]
            text = block["text"]

            if kind == "heading":
                level = min(block["level"], 3)
                doc.add_heading(text, level=level)
            elif kind == "bullet":
                para = doc.add_paragraph(style="List Bullet")
                _add_runs(para, text)
            elif kind == "numbered":
                para = doc.add_paragraph(style="List Number")
                _add_runs(para, text)
            else:
                para = doc.add_paragraph()
                _add_runs(para, text)

        buffer = io.BytesIO()
        doc.save(buffer)
        data = buffer.getvalue()

        return await _attach_file(
            __event_emitter__,
            __user__,
            data=data,
            filename=f"{_safe_filename(title)}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            max_size_mb=self.valves.max_size_mb,
        )


def _parse_blocks(markdown: str) -> list[dict]:
    """Split markdown into semantic block dicts: heading / bullet / numbered / paragraph."""
    blocks = []
    para_lines: list[str] = []

    def _flush():
        text = " ".join(para_lines).strip()
        if text:
            blocks.append({"kind": "paragraph", "text": text})
        para_lines.clear()

    for line in markdown.splitlines():
        stripped = line.strip()

        if not stripped:
            _flush()
            continue

        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            _flush()
            blocks.append({"kind": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            continue

        m = re.match(r"^[-*•]\s+(.*)", stripped)
        if m:
            _flush()
            blocks.append({"kind": "bullet", "text": m.group(1).strip()})
            continue

        m = re.match(r"^(\d+)[.)]\s+(.*)", stripped)
        if m:
            _flush()
            blocks.append({"kind": "numbered", "number": m.group(1), "text": m.group(2).strip()})
            continue

        para_lines.append(stripped)

    _flush()
    return blocks


# Matches **bold**, *italic*, `code` — bold checked first to avoid partial * matches.
_INLINE_RE = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)", re.DOTALL)


def _add_runs(para, text: str) -> None:
    """Add inline-markdown-formatted runs to a python-docx paragraph."""
    last = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > last:
            para.add_run(text[last : m.start()])
        full = m.group(0)
        if full.startswith("**"):
            para.add_run(m.group(2)).bold = True
        elif full.startswith("*"):
            para.add_run(m.group(3)).italic = True
        else:
            run = para.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def _safe_filename(name: str, fallback: str = "document") -> str:
    name = re.sub(r"[^\w\-. ]", "_", name).strip()
    return name[:120] or fallback


async def _attach_file(
    event_emitter,
    user: Optional[dict],
    *,
    data: bytes,
    filename: str,
    content_type: str,
    max_size_mb: int = 10,
) -> str:
    if event_emitter is None:
        return f"(no event emitter; would attach {filename}, {len(data)} bytes)"
    if len(data) > max_size_mb * 1024 * 1024:
        return f"Refusing to attach: {filename} is {len(data) / 1e6:.1f} MB (limit {max_size_mb} MB)."

    from open_webui.models.files import Files, FileForm
    from open_webui.storage.provider import Storage

    file_id = str(uuid.uuid4())
    user_id = (user or {}).get("id", "")
    storage_filename = f"{file_id}_{filename}"

    _, storage_path = Storage.upload_file(io.BytesIO(data), storage_filename, {})

    record = await Files.insert_new_file(
        user_id,
        FileForm(
            id=file_id,
            filename=filename,
            path=storage_path,
            meta={"name": filename, "content_type": content_type, "size": len(data)},
        ),
    )
    if record is None:
        return f"Failed to register file {filename} with Open WebUI."

    await event_emitter(
        {
            "type": "files",
            "data": {
                "files": [
                    {
                        "type": "file",
                        "id": record.id,
                        "url": f"/api/v1/files/{record.id}",
                        "name": filename,
                        "size": len(data),
                        "status": "uploaded",
                        "error": "",
                        "itemId": str(uuid.uuid4()),
                    }
                ],
            },
        }
    )

    return f"Created **{filename}** ({len(data):,} bytes). The file is now attached to the chat."
